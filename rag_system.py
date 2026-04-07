import os
from langchain_community.document_loaders import PyMuPDFLoader, TextLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_core.documents import Document

from docx import Document as DocxDocument
from rank_bm25 import BM25Okapi
from sentence_transformers import CrossEncoder


def load_docx_custom(file_path):
    doc = DocxDocument(file_path)
    text = []

    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text)

    return [Document(page_content="\n".join(text), metadata={"source": os.path.basename(file_path)})]


class HybridRAG:
    def __init__(self, persist_dir="chromadb", emb_model="sentence-transformers/all-mpnet-base-v2", k=5):
        self.k = k

        self.emb_model = HuggingFaceEmbeddings(model_name=emb_model)

        self.vectorstore = Chroma(
            persist_directory=persist_dir,
            embedding_function=self.emb_model
        )

        self.retriever = self.vectorstore.as_retriever(search_kwargs={"k": k})

        self._build_bm25()

        self.reranker = CrossEncoder('cross-encoder/ms-marco-MiniLM-L-6-v2')


    def add_documents(self, folder_path):
        docs = []

        for file in os.listdir(folder_path):
            path = os.path.join(folder_path, file)

            if file.endswith(".pdf"):
                loader = PyMuPDFLoader(path)
                docs.extend(loader.load())

            elif file.endswith(".txt"):
                loader = TextLoader(path)
                docs.extend(loader.load())

            elif file.endswith(".docx"):
                docs.extend(load_docx_custom(path))

        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        chunks = splitter.split_documents(docs)

        self.vectorstore.add_documents(chunks)
        self.vectorstore.persist()

        self._build_bm25()


    def _build_bm25(self):
        data = self.vectorstore.get()
        self.text_chunks = data["documents"]
        self.metadatas = data["metadatas"]

        # Prevent ZeroDivisionError if no documents are present
        if not self.text_chunks:
            self.bm25 = None
            return

        tokenized = [doc.lower().split() for doc in self.text_chunks]
        self.bm25 = BM25Okapi(tokenized)


    def hybrid_search(self, query):
        vector_docs = self.retriever.invoke(query)

        vector_results = [{
            "text": d.page_content,
            "meta": d.metadata,
            "score": 1.0
        } for d in vector_docs]

        # FIX: Handle case where BM25 is not built due to no documents
        if self.bm25 is None:
            return vector_results

        bm25_scores = self.bm25.get_scores(query.lower().split())
        top_idx = sorted(range(len(bm25_scores)), key=lambda i: bm25_scores[i], reverse=True)[:self.k]

        bm25_results = [{
            "text": self.text_chunks[i],
            "meta": self.metadatas[i],
            "score": bm25_scores[i]
        } for i in top_idx]

        combined = {}

        for item in vector_results + bm25_results:
            if item["text"] not in combined:
                combined[item["text"]] = item
            else:
                combined[item["text"]]["score"] += item["score"]

        return list(combined.values())


    def rerank(self, query, results):
        pairs = [(query, r["text"]) for r in results]
        scores = self.reranker.predict(pairs)

        for i, r in enumerate(results):
            r["score"] = scores[i]

        return sorted(results, key=lambda x: x["score"], reverse=True)[:self.k]


    def query(self, query, model):
        results = self.hybrid_search(query)
        results = self.rerank(query, results)

        context = "\n\n".join([r["text"] for r in results])

        prompt = f"""
You are an AI assistant answering from multiple documents.

Rules:
- Use ONLY the provided context
- Prefer the most relevant source
- Do NOT mix unrelated documents
- If partial info exists, answer carefully
- If not found, say "I don't know"
- Keep answer concise (4–7 lines max)

Context:
{}

Question:
{}

Answer:"""

        response = model.generate_content(prompt)

        sources = list(set(r["meta"].get("source", "") for r in results))

        return response.text, sources

